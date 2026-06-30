from pathlib import Path
import tempfile
from typing import Optional

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_paragraph_rtl_direction(paragraph):
    """Set actual RTL text direction (not just alignment) for a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.paragraph_format.bidi = True
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_run_rtl_direction(run):
    """Set RTL direction and complex-script for text runs"""
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    cs = OxmlElement('w:cs')
    cs.set(qn('w:val'), '1')
    rPr.append(cs)
    # ensure complex script font flag
    try:
        run.font.complex_script = True
    except Exception:
        pass


def set_document_default_rtl(doc):
    """Set document default to RTL"""
    try:
        settings = doc.settings
        element = settings.element
        defaultTabStop = OxmlElement('w:defaultTabStop')
        defaultTabStop.set(qn('w:val'), '708')
        element.append(defaultTabStop)
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        element.append(bidi)
    except Exception as e:
        print(f"Warning: Could not set document-level RTL: {e}")


def create_rtl_paragraph(doc, text="", style=None):
    """Create a paragraph with proper RTL text direction"""
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    set_paragraph_rtl_direction(para)
    if text:
        # preserve line breaks
        for line in text.splitlines():
            run = para.add_run(line)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            set_run_rtl_direction(run)
            para.add_run()  # new line
    return para


def markdown_to_docx(
    markdown_text: str,
    *,
    out_path: Optional[str | Path] = None,
    is_rtl: bool = False,
) -> str:
    """
    Convert markdown to Word with proper RTL text direction support
    """
    if out_path is None:
        tmpdir = tempfile.mkdtemp(prefix="ocr_ai_")
        out_path = Path(tmpdir) / "document.docx"
    else:
        out_path = Path(out_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    if is_rtl:
        set_document_default_rtl(doc)

    html = markdown.markdown(
        markdown_text,
        extensions=["tables", "fenced_code", "nl2br"]
    )
    soup = BeautifulSoup(html, "html.parser")
    elements = soup.find_all(['h1','h2','h3','h4','p','table','ul','ol','pre'])
    for element in elements:
        tag = element.name.lower()
        if tag.startswith('h') and tag[1].isdigit():
            level = min(int(tag[1]), 9)
            text = element.get_text(separator="\n")
            if text:
                heading = doc.add_heading(text, level=level)
                if is_rtl:
                    set_paragraph_rtl_direction(heading)
                    for run in heading.runs:
                        set_run_rtl_direction(run)
        elif tag == 'p':
            text = element.get_text(separator="\n")
            if text:
                if is_rtl:
                    create_rtl_paragraph(doc, text)
                else:
                    # split on lines to preserve breaks
                    para = doc.add_paragraph()
                    for line in text.splitlines():
                        run = para.add_run(line)
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                        para.add_run()
        elif tag == 'table':
            rows = [[cell.get_text(strip=True) for cell in tr.find_all(['td','th'])]
                    for tr in element.find_all('tr')]
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = 'Table Grid'
                if is_rtl:
                    tblPr = table._tbl.tblPr
                    bidiVisual = OxmlElement('w:bidiVisual')
                    bidiVisual.set(qn('w:val'), '1')
                    tblPr.append(bidiVisual)
                    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        para = cell.add_paragraph()
                        if is_rtl:
                            set_paragraph_rtl_direction(para)
                        run = para.add_run(cell_text)
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
                        if is_rtl:
                            set_run_rtl_direction(run)
                doc.add_paragraph()
        elif tag in ('ul','ol'):
            style = 'List Number' if tag == 'ol' else 'List Bullet'
            for li in element.find_all('li', recursive=False):
                text = li.get_text(separator="\n")
                para = doc.add_paragraph(text, style=style)
                if is_rtl:
                    set_paragraph_rtl_direction(para)
                    for run in para.runs:
                        set_run_rtl_direction(run)
        elif tag == 'pre':
            code = element.get_text()
            para = doc.add_paragraph()
            run = para.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            if is_rtl:
                set_paragraph_rtl_direction(para)
                set_run_rtl_direction(run)
    doc.save(out_path)
    return str(Path(out_path).resolve())